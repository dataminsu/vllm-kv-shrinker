"""
KV Cache Pruning Research Report Generator
- Downloads papers from arXiv
- Generates comprehensive Word report
"""

import datetime
import os
import time
import urllib.error
import urllib.request

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PAPERS_DIR = os.path.join(os.path.dirname(__file__), "papers")
os.makedirs(PAPERS_DIR, exist_ok=True)

# =============================================================================
# PAPER DATABASE
# =============================================================================
PAPERS = [
    {
        "id": "P01",
        "title": "Efficient Memory Management for Large Language Model Serving with PagedAttention",
        "authors": "Woosuk Kwon et al.",
        "venue": "SOSP 2023",
        "year": 2023,
        "arxiv": "2309.06180",
        "citations": 5316,
        "category": "Background",
        "problem": (
            "LLM 서빙 시스템에서 KV 캐시는 GPU 메모리의 최대 30%를 차지하며, "
            "기존의 연속 메모리 할당 방식은 내부/외부 단편화(fragmentation)와 "
            "요청 간 KV 캐시 공유 불가로 인해 심각한 메모리 낭비를 유발한다."
        ),
        "solution": (
            "운영체제의 가상 메모리 및 페이징 기법에서 영감을 받은 "
            "PagedAttention 알고리즘을 제안. KV 캐시를 고정 크기의 블록(page)으로 "
            "관리하며, 비연속적 메모리 블록을 logical→physical 테이블로 매핑. "
            "이를 기반으로 vLLM 서빙 시스템을 구현하여 메모리 낭비를 거의 0에 "
            "가깝게 줄이고, 요청 간/내 KV 캐시 공유를 가능하게 한다."
        ),
        "relation": (
            "본 아이디어의 구현 타겟인 vLLM의 핵심 아키텍처 논문. "
            "KVShrinker를 vLLM에 통합할 때 PagedAttention의 블록 단위 메모리 "
            "관리와 호환되어야 한다. 특히 블록 단위로 토큰을 evict/prune할 경우 "
            "PagedAttention의 block table과의 정합성 관리가 핵심 설계 과제이다."
        ),
    },
    {
        "id": "P02",
        "title": "Efficient Streaming Language Models with Attention Sinks",
        "authors": "Guangxuan Xiao et al.",
        "venue": "ICLR 2024",
        "year": 2023,
        "arxiv": "2309.17453",
        "citations": 1592,
        "category": "KV Pruning - Token Eviction",
        "problem": (
            "LLM은 학습 시 사용한 context window를 초과하는 시퀀스를 처리하면 "
            "성능이 급격히 저하된다. KV 캐시를 window 크기로 제한하면 "
            "초기 토큰이 제거될 때 perplexity가 폭발적으로 증가한다."
        ),
        "solution": (
            "초기 토큰('attention sink')이 attention score의 대부분을 흡수하는 "
            "현상을 발견. StreamingLLM은 attention sink 토큰(초기 4개)과 "
            "최근 window 토큰만 KV 캐시에 유지하는 sliding window 방식 제안. "
            "파인튜닝 없이 4M 토큰 이상의 무한 스트리밍 추론을 가능하게 함."
        ),
        "relation": (
            "본 아이디어의 RAG-aware pruning 관점에서 중요한 시사점을 제공. "
            "Attention sink 토큰은 RAG 검색 키워드와 무관하게 항상 보존해야 "
            "한다는 설계 원칙을 세울 수 있다. 즉, KVShrinker는 "
            "(1) attention sink 토큰, (2) RAG 키워드 매칭 토큰, (3) 최근 window "
            "세 가지를 보존하는 hybrid 전략으로 설계해야 한다."
        ),
    },
    {
        "id": "P03",
        "title": "H2O: Heavy-Hitter Oracle for Efficient Generative Inference of Large Language Models",
        "authors": "Zhenyu Zhang et al.",
        "venue": "NeurIPS 2023",
        "year": 2023,
        "arxiv": "2306.14048",
        "citations": 611,
        "category": "KV Pruning - Token Eviction",
        "problem": (
            "LLM 추론에서 KV 캐시 메모리가 시퀀스 길이에 비례해 증가하여 "
            "긴 시퀀스 처리 시 메모리 부족 현상이 발생한다. 어떤 토큰의 "
            "KV를 버릴지 결정하는 최적 정책이 없었다."
        ),
        "solution": (
            "누적 attention score의 합이 높은 'Heavy Hitter(H2)' 토큰이 "
            "미래 예측에 가장 중요하다는 가설을 수립하고 이론적으로 증명. "
            "H2O는 H2 토큰과 최근 토큰을 동적으로 균형 있게 유지하는 "
            "greedy eviction 알고리즘으로, 파인튜닝 없이 KV 캐시를 5배 압축."
        ),
        "relation": (
            "본 아이디어의 직접적 선행 연구. H2O는 attention score만으로 "
            "중요도를 판단하지만, 본 아이디어는 RAG 검색 시스템의 lexical 신호를 "
            "추가로 활용한다. H2O의 Heavy Hitter와 RAG 키워드가 겹치지 않는 "
            "경우(즉, RAG 키워드가 attention을 덜 받지만 의미적으로 중요한 경우) "
            "본 아이디어가 H2O 대비 우수함을 입증하는 ablation이 필요하다."
        ),
    },
    {
        "id": "P04",
        "title": "Scissorhands: Exploiting the Persistence of Importance Hypothesis for LLM KV Cache Compression at Test Time",
        "authors": "Zichang Liu et al.",
        "venue": "NeurIPS 2023",
        "year": 2023,
        "arxiv": "2305.17118",
        "citations": 359,
        "category": "KV Pruning - Token Eviction",
        "problem": (
            "KV 캐시 압축 시 어떤 토큰을 제거해야 하는지 매 스텝마다 "
            "비용이 큰 recomputation 없이 결정하는 방법이 없었다."
        ),
        "solution": (
            "'Persistence of Importance' 가설: 한 스텝에서 중요했던 pivot 토큰은 "
            "미래에도 계속 중요하다. 이 가설에 기반해 과거 attention history를 "
            "활용한 eviction 정책 ScissorHands를 제안. 파인튜닝 없이 고정 "
            "메모리 budget 내에서 KV 캐시 관리."
        ),
        "relation": (
            "RAG-aware pruning에서 중요한 인사이트: RAG 키워드는 'persistence'를 "
            "외부에서 정의하는 역할을 한다. 즉, attention history 없이도 "
            "RAG 신호를 통해 어떤 토큰이 '지속적으로 중요한지' 사전에 알 수 있다. "
            "이는 ScissorHands가 online으로 학습하는 importance를 RAG가 "
            "offline(prefill 이전)에 제공하는 방식으로 볼 수 있다."
        ),
    },
    {
        "id": "P05",
        "title": "Model Tells You What to Discard: Adaptive KV Cache Compression for LLMs",
        "authors": "Suyu Ge et al.",
        "venue": "ICLR 2024",
        "year": 2023,
        "arxiv": "2310.01801",
        "citations": 417,
        "category": "KV Pruning - Adaptive",
        "problem": (
            "기존 KV 압축 방법은 모든 attention head와 layer에 동일한 "
            "압축 전략을 적용하여 헤드별 특성을 무시하는 문제가 있었다."
        ),
        "solution": (
            "모델의 attention head를 Local(최근 토큰 선호), Strongly Local, "
            "특수 패턴 등으로 분류하고, 각 헤드 유형에 맞는 적응형 KV 압축 "
            "전략(FastGen)을 적용. 사전 프로파일링을 통해 헤드 타입을 파악하고, "
            "inference 시 타입별로 다른 eviction 정책 적용."
        ),
        "relation": (
            "본 아이디어의 RAG-aware 접근에 head-type 분류를 결합할 수 있다. "
            "Strongly Local 헤드는 최근 토큰에만 집중하므로 RAG 신호가 덜 "
            "유효하지만, semantic/global 헤드에서는 RAG 키워드 보존이 "
            "더 중요할 수 있다. Head-type별로 RAG 신호 가중치를 다르게 "
            "적용하는 확장 아이디어로 활용 가능."
        ),
    },
    {
        "id": "P06",
        "title": "SnapKV: LLM Knows What You are Looking for Before Generation",
        "authors": "Yuhong Li et al.",
        "venue": "NeurIPS 2024",
        "year": 2024,
        "arxiv": "2404.14469",
        "citations": 494,
        "category": "KV Pruning - Query-Aware",
        "problem": (
            "기존 KV 압축은 query-agnostic하여 현재 query와 무관한 토큰도 "
            "보존하는 비효율이 있었다. 긴 문서에서 특정 질문에 필요한 "
            "토큰만 선별적으로 보존하는 방법이 없었다."
        ),
        "solution": (
            "Instruction(query) 토큰의 attention pattern을 관찰하여 어떤 "
            "prefix 토큰이 중요한지 generation 이전에 미리 파악. "
            "Observation window의 attention score를 voting으로 집계하고, "
            "중요 토큰 주변에 pooling을 적용해 context 보존. "
            "파인튜닝 없이 KV 캐시를 최대 10배 이상 압축."
        ),
        "relation": (
            "본 아이디어와 가장 유사한 선행 연구. SnapKV는 query(instruction)의 "
            "attention을 통해 중요 토큰을 식별하지만, RAG 시나리오에서는 "
            "query가 항상 명확하지 않거나 multi-turn일 수 있다. "
            "본 아이디어는 RAG retriever의 lexical 신호를 추가 채널로 활용하여 "
            "SnapKV가 놓치는 경우(query와 직접 연관되지 않지만 검색된 지식에서 "
            "중요한 토큰)를 커버한다. RAG-SnapKV 결합 실험이 필요하다."
        ),
    },
    {
        "id": "P07",
        "title": "PyramidKV: Dynamic KV Cache Compression based on Pyramidal Information Funneling",
        "authors": "Zefan Cai et al.",
        "venue": "arXiv 2024",
        "year": 2024,
        "arxiv": "2406.02069",
        "citations": 226,
        "category": "KV Pruning - Layer-wise",
        "problem": (
            "모든 layer에 동일한 KV 캐시 budget을 할당하는 것은 비효율적. "
            "하위 layer는 많은 토큰에 분산 attention하고, 상위 layer는 "
            "소수의 중요 토큰에 집중 attention하는 패턴이 있다."
        ),
        "solution": (
            "Layer별로 attention entropy를 분석하여 상위 layer일수록 더 "
            "적은 KV 캐시 budget을 할당하는 피라미드 구조 제안. "
            "동일한 총 메모리 예산 내에서 flat allocation보다 "
            "유의미하게 높은 성능을 달성."
        ),
        "relation": (
            "본 아이디어의 vLLM 구현 시 layer-wise budget allocation에 "
            "PyramidKV 전략을 채택할 수 있다. RAG-aware 신호를 "
            "layer별로 다른 강도로 적용하는 방법과 결합하면, "
            "상위 layer(semantic 처리)에서 RAG 키워드를 더 강하게 "
            "보존하고 하위 layer에서는 완화하는 전략이 가능하다."
        ),
    },
    {
        "id": "P08",
        "title": "PyramidInfer: Pyramid KV Cache Compression for High-throughput LLM Inference",
        "authors": "Dongjie Yang et al.",
        "venue": "ACL 2024",
        "year": 2024,
        "arxiv": "2405.12532",
        "citations": 129,
        "category": "KV Pruning - Layer-wise",
        "problem": (
            "LLM의 throughput을 높이기 위해 KV 캐시를 줄여야 하지만, "
            "어떤 layer에서 어떤 토큰을 제거할지 결정하는 원칙이 불명확했다."
        ),
        "solution": (
            "Layer가 깊어질수록 미래 생성에 영향을 미치는 중요 key의 수가 "
            "감소한다는 점을 attention consistency로 측정. 상위 layer에서 "
            "더 적은 KV를 유지하는 피라미드 구조로 throughput을 크게 향상."
        ),
        "relation": (
            "PyramidKV와 유사하나 throughput 관점에서 접근. "
            "본 아이디어의 vLLM 구현에서 처리량 목표와 RAG-aware 정확도 "
            "목표를 동시에 달성하기 위한 레퍼런스로 활용. "
            "특히 layer consistency metric이 RAG 키워드 토큰의 "
            "'persistence score'를 결합할 수 있는 후보."
        ),
    },
    {
        "id": "P09",
        "title": "KIVI: A Tuning-Free Asymmetric 2bit Quantization for KV Cache",
        "authors": "Zirui Liu et al.",
        "venue": "ICML 2024",
        "year": 2024,
        "arxiv": "2402.02750",
        "citations": 408,
        "category": "KV Quantization (Comparison)",
        "problem": (
            "KV 캐시의 메모리 사용량을 줄이면서 생성 품질을 유지하는 "
            "양자화 방법이 필요했다. Key와 Value의 통계적 특성이 달라 "
            "동일한 양자화 전략 적용 시 오차가 크다."
        ),
        "solution": (
            "Key 캐시는 채널 차원으로 per-channel 양자화, Value 캐시는 "
            "토큰 차원으로 per-token 양자화하는 비대칭 2-bit 양자화 방법. "
            "소수의 residual full-precision 토큰을 유지하여 오차 보상. "
            "파인튜닝 없이 2-bit 압축에서도 16-bit에 근접한 성능."
        ),
        "relation": (
            "[Quantization vs. Pruning 비교 대상] KIVI는 모든 토큰의 KV를 "
            "낮은 정밀도로 보존하지만, 본 아이디어(pruning 기반)는 "
            "선택된 토큰의 KV를 full precision으로 보존하고 나머지를 제거. "
            "RAG 환경에서 중요 키워드 토큰은 full precision이 필수적이므로 "
            "pruning 접근이 quantization보다 유리할 수 있다. "
            "KIVI는 하드웨어 friendly하나, pruning은 semantic 제어가 용이하다."
        ),
    },
    {
        "id": "P10",
        "title": "GEAR: An Efficient KV Cache Compression Recipe for Near-Lossless Generative Inference of LLM",
        "authors": "Hao Kang et al.",
        "venue": "arXiv 2024",
        "year": 2024,
        "arxiv": "2403.05527",
        "citations": 144,
        "category": "KV Quantization (Comparison)",
        "problem": (
            "단순 양자화는 outlier 토큰에서 큰 오차를 유발한다. "
            "특히 이상값 토큰(outlier)의 KV가 양자화되면 "
            "attention 패턴이 크게 왜곡되어 성능 저하가 발생한다."
        ),
        "solution": (
            "KV 캐시를 quantization + low-rank approximation + sparse matrix로 "
            "분해하는 GEAR 프레임워크. 양자화 오차를 low-rank 행렬로 보상하고, "
            "남은 outlier는 sparse하게 full precision 유지. "
            "2.38배 throughput 향상, 2.29배 메모리 절감 달성."
        ),
        "relation": (
            "[Quantization vs. Pruning 비교 대상] GEAR는 압축 품질이 높지만 "
            "low-rank 분해와 sparse 행렬 관리의 복잡도가 크다. "
            "본 RAG-aware pruning은 external knowledge(RAG 신호)를 활용한 "
            "선택적 보존으로, GEAR보다 단순하면서도 RAG 시나리오에서 "
            "더 높은 정확도를 달성할 수 있다는 가설을 실험으로 검증해야 한다."
        ),
    },
    {
        "id": "P11",
        "title": "Get More with LESS: Synthesizing Recurrence with KV Cache Compression for Efficient LLM Inference",
        "authors": "Harry Dong et al.",
        "venue": "ICML 2024",
        "year": 2024,
        "arxiv": "2402.09398",
        "citations": 83,
        "category": "KV Pruning - Recurrence",
        "problem": (
            "KV 캐시 eviction은 제거된 토큰의 정보를 완전히 잃는다. "
            "Eviction된 토큰의 정보를 어떻게 보존하면서도 메모리를 줄일 수 있는가?"
        ),
        "solution": (
            "KV 캐시 eviction과 recurrence 메커니즘을 결합한 LESS 프레임워크. "
            "제거된 KV를 compact한 recurrent state로 압축하여 정보 손실을 "
            "최소화. Eviction + summary state의 두 트랙으로 메모리를 줄이면서 "
            "long-range dependency를 유지."
        ),
        "relation": (
            "본 아이디어에서 RAG 문서의 비-키워드 토큰을 단순 제거하는 대신, "
            "recurrent state로 압축하는 hybrid 방법을 제안할 수 있다. "
            "즉, RAG 키워드 토큰은 full KV 보존, 나머지는 LESS 방식의 "
            "recurrent 압축을 적용하는 2-tier 전략."
        ),
    },
    {
        "id": "P12",
        "title": "ThinK: Thinner Key Cache by Query-Driven Pruning",
        "authors": "Yuhui Xu et al.",
        "venue": "ICLR 2025",
        "year": 2024,
        "arxiv": "2407.21018",
        "citations": 49,
        "category": "KV Pruning - Dimension Pruning",
        "problem": (
            "기존 KV 캐시 압축은 토큰 단위(token-level)로 KV를 evict하지만, "
            "key 텐서의 채널(dimension) 차원에도 중복성이 존재한다."
        ),
        "solution": (
            "Key 캐시의 채널 차원을 query와의 상관관계 기반으로 선택적으로 "
            "pruning하는 ThinK 방법. Query-driven 방식으로 중요하지 않은 "
            "key 채널을 제거하여 메모리를 줄이면서도 attention 품질 유지. "
            "토큰 삭제 없이 채널 차원 압축만으로 20-30% 메모리 절감."
        ),
        "relation": (
            "Token-level pruning과 channel-level pruning은 orthogonal하다. "
            "본 아이디어의 RAG-aware token-level pruning과 ThinK의 "
            "channel-level pruning을 결합하면 두 방향으로 동시에 압축하는 "
            "2D KV compression이 가능하다. 이는 NeurIPS 논문의 contribution이 "
            "될 수 있는 아이디어다."
        ),
    },
    {
        "id": "P13",
        "title": "CaM: Cache Merging for Memory-efficient LLMs Inference",
        "authors": "Yuxin Zhang et al.",
        "venue": "ICML 2024",
        "year": 2024,
        "arxiv": None,  # No ArXiv found
        "citations": 63,
        "category": "KV Pruning - Merging",
        "problem": (
            "단순 eviction은 정보를 완전히 손실한다. 제거 대상 토큰의 "
            "정보를 보존 토큰에 합산(merge)하는 방법이 없었다."
        ),
        "solution": (
            "제거할 토큰의 KV를 가장 유사한 보존 토큰의 KV에 merge하여 "
            "정보 손실을 최소화하는 CaM 방법. Eviction policy와 merge policy를 "
            "분리하여 다양한 base eviction 방법에 플러그인 형태로 적용 가능."
        ),
        "relation": (
            "본 아이디어에서 RAG 비-키워드 토큰을 evict할 때 단순 제거 대신 "
            "가장 가까운 RAG 키워드 토큰에 merge하는 방식을 적용할 수 있다. "
            "이는 정보 손실을 줄이면서도 RAG 키워드 토큰의 중요성을 강화하는 "
            "시너지 효과를 낼 수 있다."
        ),
    },
    {
        "id": "P14",
        "title": "ClusterKV: Manipulating LLM KV Cache in Semantic Space for Recallable Compression",
        "authors": "Guangda Liu et al.",
        "venue": "DAC 2024",
        "year": 2024,
        "arxiv": "2412.03213",
        "citations": 48,
        "category": "KV Pruning - Semantic",
        "problem": (
            "토큰을 개별적으로 evict하면 의미적으로 관련된 토큰들이 "
            "분리되어 semantic coherence가 손상된다. "
            "또한 evict된 토큰을 나중에 recall하는 메커니즘이 없었다."
        ),
        "solution": (
            "KV 캐시를 semantic space에서 클러스터링하여 관련 토큰들을 "
            "함께 보존/제거하는 ClusterKV. Query와 의미적으로 유사한 "
            "클러스터는 보존하고, 무관한 클러스터는 압축 저장 후 필요 시 recall."
        ),
        "relation": (
            "본 아이디어의 RAG-aware 신호와 semantic clustering을 결합할 수 있다. "
            "RAG 키워드와 semantic하게 유사한 클러스터를 우선 보존하면, "
            "단순 lexical matching을 넘어 의미적으로 관련된 토큰까지 "
            "보호하는 더 robust한 방법이 된다. 이는 dense retrieval과도 "
            "자연스럽게 연결된다."
        ),
    },
    {
        "id": "P15",
        "title": "D2O: Dynamic Discriminative Operations for Efficient Long-Context Inference of Large Language Models",
        "authors": "Zhongwei Wan et al.",
        "venue": "ICLR 2025",
        "year": 2024,
        "arxiv": "2406.13035",
        "citations": 13,
        "category": "KV Pruning - Dynamic",
        "problem": (
            "Long-context 추론에서 KV 캐시가 폭발적으로 증가하는 문제. "
            "단일 기준(예: attention score)으로는 모든 layer/head에서 "
            "최적의 eviction을 보장하기 어렵다."
        ),
        "solution": (
            "각 layer와 head의 특성에 따라 보존(Retain) 또는 제거(Dismiss) "
            "정책을 동적으로 선택하는 D2O 방법. Layer별 diversity를 측정하여 "
            "adaptive하게 KV 캐시 크기를 조절."
        ),
        "relation": (
            "본 아이디어에서 RAG 신호의 강도(retrieval score)에 따라 "
            "dynamic하게 KV budget을 조절하는 방법과 결합 가능. "
            "High retrieval score → 더 많은 KV 보존, Low score → 더 공격적 pruning. "
            "D2O의 dynamic 프레임워크를 RAG score guided budget allocation에 적용."
        ),
    },
    {
        "id": "P16",
        "title": "LoRC: Low-Rank Compression for LLMs KV Cache with a Progressive Compression Strategy",
        "authors": "Rongzhi Zhang et al.",
        "venue": "arXiv 2024",
        "year": 2024,
        "arxiv": "2410.03111",
        "citations": 34,
        "category": "KV Pruning - Low-rank",
        "problem": (
            "KV 캐시를 단순 eviction하면 정보가 소실되고, "
            "quantization은 하드웨어 지원이 필요하다. "
            "행렬 분해(matrix decomposition) 기반의 손실 최소화 압축이 필요."
        ),
        "solution": (
            "KV 캐시 행렬을 low-rank 근사로 압축하는 LoRC. "
            "Progressive 전략으로 점진적으로 rank를 줄여나가며, "
            "attention 패턴 보존을 목표로 SVD 기반 압축 적용."
        ),
        "relation": (
            "Low-rank 방법은 RAG-aware pruning과 직교적으로 결합 가능. "
            "RAG 키워드 토큰에 대응하는 KV는 full-rank 보존, "
            "비-키워드 토큰은 low-rank 압축을 적용하는 mixed-rank 전략. "
            "이는 GEAR의 mixed-precision과 유사하나 RAG 신호 기반이라는 점이 차별점."
        ),
    },
    {
        "id": "P17",
        "title": "ZigZagKV: Dynamic KV Cache Compression for Long-context Modeling based on Layer Uncertainty",
        "authors": "M. Zhong et al.",
        "venue": "COLING 2025",
        "year": 2024,
        "arxiv": "2412.09036",
        "citations": 5,
        "category": "KV Pruning - Layer-wise",
        "problem": (
            "Layer마다 attention 분포의 불확실성(uncertainty)이 다르지만 "
            "기존 방법들은 이를 무시하고 일률적 budget을 적용한다."
        ),
        "solution": (
            "각 layer의 attention entropy를 기반으로 uncertainty를 측정하고, "
            "uncertainty가 높은 layer(더 많은 토큰에 분산 attention)에는 "
            "더 많은 KV budget을 할당하는 ZigZag 패턴의 dynamic budget 전략."
        ),
        "relation": (
            "RAG 시나리오에서 retrieved document 처리 시 특정 layer의 "
            "uncertainty가 증가할 수 있다. RAG 문서의 entropy 패턴과 "
            "ZigZagKV의 layer uncertainty를 결합하면, RAG 내용이 "
            "높은 uncertainty를 유발하는 layer에서 RAG 키워드를 더 강하게 보존하는 "
            "adaptive 전략을 설계할 수 있다."
        ),
    },
    {
        "id": "P18",
        "title": "DiffKV: Differentiated Memory Management for Large Language Models with Parallel KV Compaction",
        "authors": "Yanqi Zhang et al.",
        "venue": "SOSP 2024",
        "year": 2024,
        "arxiv": "2412.03131",
        "citations": 7,
        "category": "KV Pruning - System",
        "problem": (
            "Key와 Value 캐시의 attention에 미치는 영향이 다르고, "
            "토큰 중요도와 head별 sparsity 패턴도 다양하다. "
            "이를 통합적으로 고려한 시스템 수준의 KV 관리 방법이 없었다."
        ),
        "solution": (
            "Key-Value 비대칭성, 토큰 중요도, head별 sparsity를 동시에 고려한 "
            "DiffKV 프레임워크. Parallel KV compaction으로 압축 오버헤드를 줄이며 "
            "세 가지 differentiation을 활용한 최적 KV 관리."
        ),
        "relation": (
            "vLLM 기반 구현 관점에서 참고할 시스템 설계. "
            "본 아이디어의 KVShrinker 구현 시 Key와 Value를 비대칭적으로 처리하는 "
            "DiffKV의 접근법을 채택 가능. 특히 RAG 키워드에 해당하는 Key는 "
            "full precision/dimension 보존, Value는 부분 압축하는 전략과 결합."
        ),
    },
    {
        "id": "P19",
        "title": "KVzip: Query-Agnostic KV Cache Compression with Context Reconstruction",
        "authors": "Jang-Hyun Kim et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "arxiv": "2505.23416",
        "citations": 20,
        "category": "KV Pruning - Query-Agnostic",
        "problem": (
            "Query-aware 방법(SnapKV 등)은 query를 알아야 하므로 prefill 단계에서 "
            "미리 KV를 압축할 수 없다. Multi-turn 또는 prefix caching 환경에서 "
            "query-agnostic한 압축이 필요하다."
        ),
        "solution": (
            "Context reconstruction loss를 최소화하는 방향으로 "
            "query-agnostic KV 압축을 수행하는 KVzip. "
            "어떤 query가 들어와도 context를 잘 재구성할 수 있는 KV subset 선택."
        ),
        "relation": (
            "RAG 시나리오에서도 일부는 query-agnostic 압축이 필요하다. "
            "예를 들어 shared prefix(시스템 프롬프트)나 document 압축은 "
            "query-agnostic. 본 아이디어는 RAG 키워드가 있을 때 "
            "query-aware 신호로 KVzip을 보강하는 방향으로 확장 가능. "
            "RAG-aware와 query-agnostic의 hybrid 접근법이 될 수 있다."
        ),
    },
    {
        "id": "P20",
        "title": "SparK: Query-Aware Unstructured Sparsity with Recoverable KV Cache Channel Pruning",
        "authors": "Huanxuan Liao et al.",
        "venue": "AAAI 2025",
        "year": 2025,
        "arxiv": "2508.15212",
        "citations": 3,
        "category": "KV Pruning - Structured Sparsity",
        "problem": (
            "구조적 KV 압축(채널 pruning)은 hardware efficiency가 높지만 "
            "정보 손실이 크고, 비구조적 압축은 그 반대다. "
            "Query-aware하면서도 recoverable한 구조적 압축이 필요하다."
        ),
        "solution": (
            "Query 정보를 활용하여 KV 캐시의 channel(dimension) 단위로 "
            "unstructured sparsity를 적용하는 SparK. "
            "Pruned channel의 정보를 복구 가능한 형태로 저장하여 "
            "필요 시 recall. Hardware-friendly 구조적 압축과 "
            "정보 보존을 동시에 달성."
        ),
        "relation": (
            "ThinK와 함께 channel-level pruning의 최신 방향. "
            "본 아이디어는 token-level에서 RAG 신호를 활용하지만, "
            "SparK의 channel-level recovery 메커니즘을 "
            "RAG 키워드 토큰의 KV channel을 선택적으로 복구하는 데 활용 가능. "
            "즉, compressed channel 중 RAG 키워드 관련 채널만 우선 복구하는 전략."
        ),
    },
    {
        "id": "P21",
        "title": "Lethe: Layer- and Time-Adaptive KV Cache Pruning for Reasoning-Intensive LLM Serving",
        "authors": "Hui Zeng et al.",
        "venue": "AAAI 2025",
        "year": 2025,
        "arxiv": "2511.06029",
        "citations": 2,
        "category": "KV Pruning - Adaptive",
        "problem": (
            "추론 집약적(reasoning-intensive) 작업에서 KV 캐시 pruning이 "
            "chain-of-thought 품질을 크게 저하시킨다. "
            "Layer와 time step에 따라 최적 pruning 전략이 달라진다."
        ),
        "solution": (
            "Layer와 time step(decoding position)을 동시에 고려하는 "
            "adaptive pruning 전략 Lethe. 초기 decoding 단계에서는 "
            "보수적 pruning, 후기 단계에서는 공격적 pruning을 적용. "
            "Reasoning 작업에서 기존 방법 대비 성능 하락 최소화."
        ),
        "relation": (
            "RAG 기반 추론에서도 reasoning 품질 보존이 중요하다. "
            "본 아이디어에서 RAG 키워드를 활용한 pruning은 특히 "
            "multi-hop reasoning이나 복잡한 QA에서 중요 토큰 보호가 핵심. "
            "Lethe의 time-adaptive 전략과 RAG-aware 신호를 결합하면 "
            "reasoning task에서 더 안정적인 성능을 기대할 수 있다."
        ),
    },
    {
        "id": "P22",
        "title": "OBCache: Optimal Brain KV Cache Pruning for Efficient Long-Context LLM Inference",
        "authors": "Yuzhe Gu et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "arxiv": "2510.07651",
        "citations": 1,
        "category": "KV Pruning - Structured",
        "problem": (
            "KV 캐시 pruning 시 어떤 토큰을 제거할지 결정하는 것이 "
            "최적 신경망 가중치 pruning과 유사한 구조적 최적화 문제임을 인식. "
            "기존 heuristic 기반 방법들의 이론적 근거가 부족하다."
        ),
        "solution": (
            "Optimal Brain Surgeon(OBS) 프레임워크를 KV 캐시 pruning에 적용한 OBCache. "
            "2차 Taylor 전개를 통해 각 KV를 제거했을 때 attention 출력에 미치는 "
            "영향을 정확히 추정하여 최적의 pruning 결정."
        ),
        "relation": (
            "이론적으로 가장 엄밀한 접근법. 본 아이디어의 RAG-aware pruning을 "
            "OBCache의 이론적 프레임워크와 결합할 수 있다. "
            "즉, RAG 키워드에 해당하는 토큰의 'pruning cost'에 페널티를 추가하여 "
            "이론적으로 최적이면서도 RAG-aware한 pruning 결정을 도출. "
            "이는 NeurIPS 수준의 이론적 기여가 될 수 있다."
        ),
    },
    {
        "id": "P23",
        "title": "HyperRAG: Enhancing Quality-Efficiency Tradeoffs in Retrieval-Augmented Generation with Reranker KV-Cache Reuse",
        "authors": "Yuwei An et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "arxiv": "2504.02921",
        "citations": 7,
        "category": "RAG + KV Cache",
        "problem": (
            "RAG 파이프라인에서 reranker와 LLM generator가 동일한 passage를 "
            "반복 처리하는 비효율이 있다. reranker의 KV 캐시를 generator에 "
            "재사용하는 방법이 없었다."
        ),
        "solution": (
            "Reranker가 생성한 KV 캐시를 generator LLM에서 직접 재사용하는 "
            "HyperRAG 프레임워크. Passage당 KV 계산을 한 번만 수행하여 "
            "latency를 크게 줄이면서 품질은 유지."
        ),
        "relation": (
            "RAG와 KV 캐시를 연결하는 가장 직접적인 선행 연구. "
            "단, HyperRAG는 KV 재사용에 초점을 맞추지만, 본 아이디어는 "
            "RAG 신호를 KV pruning 결정에 활용한다는 점에서 차별화. "
            "두 아이디어를 결합하면: RAG reranker의 lexical score를 이용해 "
            "KV pruning을 수행하면서 동시에 reranker KV를 재사용하는 "
            "통합 RAG-KV 시스템이 가능하다."
        ),
    },
    {
        "id": "P24",
        "title": "Identify Critical KV Cache in LLM Inference from an Output Perturbation Perspective",
        "authors": "Yuan Feng et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "arxiv": "2502.03805",
        "citations": 15,
        "category": "KV Pruning - Importance Scoring",
        "problem": (
            "KV 캐시의 중요도를 attention score로만 측정하면 "
            "실제 출력에 미치는 영향을 정확히 반영하지 못한다. "
            "어떤 KV가 LLM 출력에 가장 큰 영향을 미치는지 직접 측정이 필요하다."
        ),
        "solution": (
            "특정 KV를 제거했을 때 출력 분포의 perturbation(변화량)을 측정하여 "
            "중요도를 정의하는 방법. Perturbation 크기가 작은 KV를 제거하는 "
            "output-perturbation guided eviction 정책 제안."
        ),
        "relation": (
            "본 아이디어의 RAG-aware importance score와 output perturbation score를 "
            "결합하는 앙상블 방법을 제안할 수 있다. "
            "RAG 키워드 토큰의 perturbation score가 낮더라도 RAG 신호에 의해 "
            "보존하는 전략이 언제 유효한지 분석하는 것이 "
            "논문의 핵심 실험 설계가 될 수 있다."
        ),
    },
    {
        "id": "P25",
        "title": "GraphKV: Breaking the Static Selection Paradigm with Graph-Based KV Cache Eviction",
        "authors": "Xuelin Li et al.",
        "venue": "EMNLP 2025",
        "year": 2025,
        "arxiv": "2509.00388",
        "citations": 0,
        "category": "KV Pruning - Graph-based",
        "problem": (
            "기존 KV eviction은 각 토큰의 중요도를 독립적으로 평가하여 "
            "토큰 간의 상관관계(dependency)를 무시한다. "
            "서로 의존하는 토큰 집합이 분리되면 의미 손실이 크다."
        ),
        "solution": (
            "토큰 간의 attention 관계를 그래프로 모델링하고, "
            "graph-based 중요도 전파로 eviction 결정을 내리는 GraphKV. "
            "PageRank류 알고리즘으로 토큰의 구조적 중요도를 계산."
        ),
        "relation": (
            "RAG 키워드 토큰을 그래프의 '시드 노드'로 설정하고, "
            "그래프 전파를 통해 키워드와 강하게 연결된 토큰들을 함께 보존하는 "
            "RAG-seeded graph pruning 방법으로 확장 가능. "
            "이는 단순 lexical matching을 넘어 syntactic/semantic dependency를 "
            "고려한 더 정교한 RAG-aware pruning이 된다."
        ),
    },
    {
        "id": "P26",
        "title": "CompressKV: Semantic Retrieval Heads Know What Tokens are Not Important Before Generation",
        "authors": "Xiaolin Lin et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "arxiv": "2508.02401",
        "citations": 2,
        "category": "KV Pruning - Retrieval Head",
        "problem": (
            "LLM 내부의 'retrieval head'(정보를 검색하는 역할을 하는 attention head)가 "
            "실제로 어떤 토큰이 generation에 중요한지 사전에 알 수 있는지, "
            "그리고 이를 KV 압축에 활용할 수 있는지 연구가 부족하다."
        ),
        "solution": (
            "LLM 내 semantic retrieval head의 attention pattern을 분석하여 "
            "generation 이전에 중요 토큰을 식별하고 KV를 압축하는 CompressKV. "
            "Retrieval head의 attention이 높은 토큰을 우선 보존."
        ),
        "relation": (
            "본 아이디어와 매우 유사한 'retrieval' 관점을 공유하지만, "
            "CompressKV는 모델 내부의 retrieval head를 사용하는 반면 "
            "본 아이디어는 외부 RAG retriever의 신호를 사용한다. "
            "두 신호(내부 retrieval head + 외부 RAG 신호)를 결합하면 "
            "더 강력한 hybrid 방법이 될 수 있다. 이는 명시적 차별화 포인트."
        ),
    },
    {
        "id": "P27",
        "title": "When Attention Sink Emerges in Language Models: An Empirical View",
        "authors": "Xiangming Gu et al.",
        "venue": "ICLR 2025",
        "year": 2024,
        "arxiv": "2410.10781",
        "citations": 124,
        "category": "Analysis",
        "problem": (
            "StreamingLLM에서 발견된 attention sink 현상이 왜 발생하는지, "
            "어떤 조건에서 나타나는지에 대한 이론적/empirical 분석이 부족하다."
        ),
        "solution": (
            "다양한 모델과 입력에서 attention sink를 체계적으로 분석. "
            "Softmax normalization으로 인한 토큰 간 attention score 경쟁이 "
            "sink 현상의 주요 원인임을 밝힘. 모델 크기, 학습 방법, "
            "입력 특성과 attention sink의 관계를 정량적으로 분석."
        ),
        "relation": (
            "RAG-aware KV pruning 설계 시 attention sink 현상을 반드시 고려해야 한다. "
            "RAG 문서의 첫 토큰이나 특수 토큰(BOS, delimiter)이 sink 역할을 할 수 있어 "
            "이들을 RAG 키워드와 무관하게 보존해야 한다. "
            "또한 RAG 키워드 자체가 sink 토큰이 되는 경우와 아닌 경우를 "
            "구별하는 분석이 필요하다."
        ),
    },
]

# ArXiv IDs for download
ARXIV_IDS = [
    ("2309.06180", "P01_vLLM_PagedAttention.pdf"),
    ("2309.17453", "P02_StreamingLLM.pdf"),
    ("2306.14048", "P03_H2O.pdf"),
    ("2305.17118", "P04_Scissorhands.pdf"),
    ("2310.01801", "P05_FastGen.pdf"),
    ("2404.14469", "P06_SnapKV.pdf"),
    ("2406.02069", "P07_PyramidKV.pdf"),
    ("2405.12532", "P08_PyramidInfer.pdf"),
    ("2402.02750", "P09_KIVI.pdf"),
    ("2403.05527", "P10_GEAR.pdf"),
    ("2402.09398", "P11_LESS.pdf"),
    ("2407.21018", "P12_ThinK.pdf"),
    ("2412.03213", "P14_ClusterKV.pdf"),
    ("2410.03111", "P16_LoRC.pdf"),
    ("2412.09036", "P17_ZigZagKV.pdf"),
    ("2412.03131", "P18_DiffKV.pdf"),
    ("2505.23416", "P19_KVzip.pdf"),
    ("2508.15212", "P20_SparK.pdf"),
    ("2511.06029", "P21_Lethe.pdf"),
    ("2510.07651", "P22_OBCache.pdf"),
    ("2504.02921", "P23_HyperRAG.pdf"),
    ("2502.03805", "P24_CriticalKV.pdf"),
    ("2509.00388", "P25_GraphKV.pdf"),
    ("2508.02401", "P26_CompressKV.pdf"),
    ("2410.10781", "P27_AttentionSink.pdf"),
    ("2406.13035", "P15_D2O.pdf"),
]


# =============================================================================
# DOWNLOAD PAPERS
# =============================================================================
def download_papers():
    print("\n[1/2] 논문 PDF 다운로드 시작...")
    success, fail = 0, 0
    for arxiv_id, filename in ARXIV_IDS:
        dest = os.path.join(PAPERS_DIR, filename)
        if os.path.exists(dest) and os.path.getsize(dest) > 10000:
            print(f"  [SKIP] {filename} (이미 존재)")
            success += 1
            continue
        url = f"https://arxiv.org/pdf/{arxiv_id}"
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=30) as resp, open(dest, "wb") as f:
                f.write(resp.read())
            size_kb = os.path.getsize(dest) // 1024
            print(f"  [OK]   {filename} ({size_kb} KB)")
            success += 1
        except Exception as e:
            print(f"  [FAIL] {filename}: {e}")
            fail += 1
        time.sleep(1.2)  # arxiv rate limit
    print(f"  완료: {success}개 성공, {fail}개 실패\n")


# =============================================================================
# WORD DOCUMENT HELPERS
# =============================================================================
def set_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_table_row(table, cells_data, bold_first=False):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        cell.text = data
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# =============================================================================
# GENERATE REPORT
# =============================================================================
def generate_report():
    print("[2/2] Word 보고서 생성 중...")
    doc = Document()

    # --- 기본 스타일 설정 ---
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)

    # ==========================================================================
    # 표지
    # ==========================================================================
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("KV Cache Pruning 연구 동향 및\nRAG-aware KV Pruning 아이디어 평가 보고서")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        f"작성일: {datetime.datetime.now().strftime('%Y년 %m월 %d일')}\n"
        f"논문 수집 범위: 2022–2025년 (27편)\n"
        f"타겟 구현: vLLM KVShrinker 레이어\n"
        f"타겟 논문: NeurIPS / EMNLP"
    )
    sub_run.font.size = Pt(12)

    doc.add_page_break()

    # ==========================================================================
    # 목차
    # ==========================================================================
    set_heading(doc, "목 차", 1)
    toc_items = [
        "1. 연구 배경 및 목적",
        "2. KV Cache Pruning 기술 개요",
        "3. 수집 논문 분석 (27편)",
        "4. KV Quantization과 KV Pruning 비교",
        "5. 나의 아이디어 평가: RAG-aware KV Pruning in vLLM",
        "6. 미개척 연구 방향 및 Top-tier 논문 아이디어",
        "7. 실험 방법론",
        "8. 결론 및 향후 계획",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ==========================================================================
    # 1. 연구 배경
    # ==========================================================================
    set_heading(doc, "1. 연구 배경 및 목적", 1)
    doc.add_paragraph(
        "대형 언어 모델(LLM)의 추론 과정에서 KV(Key-Value) 캐시는 GPU 메모리의 "
        "최대 30–70%를 차지하며, 시퀀스 길이에 비례하여 O(n) 또는 O(n²)으로 "
        "증가한다. 이는 배치 크기 제한, 긴 문맥 처리 불가, 높은 서비스 비용이라는 "
        "세 가지 핵심 병목을 유발한다."
    )
    doc.add_paragraph(
        "본 보고서는 KV 캐시를 줄이는 핵심 기법인 'KV Pruning(Token Eviction)'을 "
        "중심으로 최근 3년간(2022–2025)의 연구를 체계적으로 정리하고, "
        "제안 아이디어인 RAG-aware KV Pruning의 실현 가능성과 연구 차별성을 평가한다."
    )
    doc.add_paragraph(
        "구현 목표: vLLM 내에 KVShrinker 레이어/객체를 추가하여 서비스 품질(generation quality) "
        "저하 없이 KV 캐시 메모리를 50–80% 절감하고, RAG 파이프라인에서 검색 신호를 "
        "활용한 semantically-aware pruning으로 기존 방법 대비 우수한 성능 달성."
    )

    # ==========================================================================
    # 2. KV Cache Pruning 기술 개요
    # ==========================================================================
    set_heading(doc, "2. KV Cache Pruning 기술 개요", 1)

    set_heading(doc, "2.1 KV 캐시와 메모리 문제", 2)
    doc.add_paragraph(
        "Transformer 기반 LLM의 autoregressive decoding에서 각 레이어의 "
        "Multi-Head Attention은 과거 모든 토큰의 Key와 Value 벡터를 저장한다. "
        "예를 들어 Llama-3 70B 모델에서 4096 토큰 시퀀스를 처리하면 "
        "KV 캐시 메모리는 약 40GB에 달하며, 이는 모델 가중치(140GB)의 약 29%에 해당한다."
    )

    set_heading(doc, "2.2 KV Pruning 방법론 분류", 2)
    doc.add_paragraph("KV Pruning 방법론은 다음과 같이 분류된다:")

    categories_table = doc.add_table(rows=1, cols=3)
    categories_table.style = 'Table Grid'
    hdr = categories_table.rows[0].cells
    for cell, text in zip(hdr, ["분류", "핵심 아이디어", "대표 논문"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
        shade_cell(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    cat_data = [
        ("Token Eviction\n(Attention Score 기반)", "누적 attention score가 낮은 토큰 제거", "H2O, Scissorhands, StreamingLLM"),
        ("Query-Aware Eviction", "현재 query의 attention pattern으로 중요 토큰 식별", "SnapKV, ThinK, SparK"),
        ("Layer-wise Adaptive", "Layer별로 다른 budget 할당", "PyramidKV, PyramidInfer, ZigZagKV"),
        ("Semantic/Clustering", "의미적 유사도 기반 토큰 그룹핑", "ClusterKV, GraphKV"),
        ("Merging/Recurrence", "제거 토큰 정보를 압축 보존", "CaM, LESS, KVReviver"),
        ("Dimension Pruning", "Token이 아닌 채널(dimension) 단위 압축", "ThinK, SparK"),
        ("System-level", "vLLM 등 serving system과의 통합", "DiffKV, PagedEviction, vLLM"),
    ]
    for row_data in cat_data:
        add_table_row(categories_table, row_data)

    doc.add_paragraph()

    set_heading(doc, "2.3 핵심 공통 관찰 사항", 2)
    key_findings = [
        "초기 토큰(BOS, delimiter)은 attention sink 역할을 하여 항상 보존해야 한다.",
        "상위 Layer일수록 소수의 중요 토큰에 집중 attention하는 경향이 있다.",
        "누적 attention score와 실제 출력 기여도(output perturbation)는 항상 일치하지 않는다.",
        "Query-aware 방법이 query-agnostic 방법보다 일반적으로 우수하다.",
        "Token eviction과 dimension pruning은 orthogonal하여 결합 가능하다.",
    ]
    for finding in key_findings:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(finding)

    doc.add_page_break()

    # ==========================================================================
    # 3. 수집 논문 분석
    # ==========================================================================
    set_heading(doc, "3. 수집 논문 분석 (27편)", 1)
    doc.add_paragraph(
        f"아래는 2022–2025년 사이 발표된 KV 캐시 압축 관련 핵심 논문 {len(PAPERS)}편의 "
        "상세 분석이다. 각 논문에 대해 (1) 해결하려는 문제, (2) 제안 방법, "
        "(3) 본 아이디어(RAG-aware KV Pruning)와의 관련성을 서술한다."
    )

    # Overview table
    overview_table = doc.add_table(rows=1, cols=5)
    overview_table.style = 'Table Grid'
    for cell, text in zip(overview_table.rows[0].cells, ["ID", "논문 제목 (약칭)", "게재지", "연도", "인용수"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for p in PAPERS:
        short_title = p["title"].split(":")[0][:50]
        add_table_row(overview_table, [p["id"], short_title, p["venue"], str(p["year"]), str(p["citations"])])

    doc.add_paragraph()

    # Detailed analysis per paper
    for paper in PAPERS:
        set_heading(doc, f"{paper['id']}. {paper['title']}", 2)

        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Table Grid'
        meta_rows = [
            ("저자", paper["authors"]),
            ("게재지 / 연도", f"{paper['venue']} ({paper['year']})"),
            ("인용 수", f"{paper['citations']:,}"),
            ("분류", paper["category"]),
        ]
        for row, (label, value) in zip(meta_table.rows, meta_rows):
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            shade_cell(row.cells[0], "D6E4F7")
            row.cells[1].text = value

        doc.add_paragraph()

        for section_title, content in [
            ("[문제]", paper["problem"]),
            ("[해결 방법]", paper["solution"]),
            ("[나의 아이디어와의 관련성]", paper["relation"]),
        ]:
            p = doc.add_paragraph()
            run_title = p.add_run(section_title + " ")
            run_title.bold = True
            run_title.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            p.add_run(content)

        doc.add_paragraph()

    doc.add_page_break()

    # ==========================================================================
    # 4. KV Quantization vs Pruning
    # ==========================================================================
    set_heading(doc, "4. KV Quantization과 KV Pruning 비교", 1)

    doc.add_paragraph(
        "KV 캐시 압축의 두 주요 패러다임인 Quantization과 Pruning은 "
        "서로 다른 trade-off를 가진다. 본 섹션은 두 방법의 강점과 약점, "
        "그리고 Pruning이 RAG 환경에서 갖는 실효적 우위를 분석한다."
    )

    comp_table = doc.add_table(rows=1, cols=3)
    comp_table.style = 'Table Grid'
    for cell, text in zip(comp_table.rows[0].cells, ["비교 항목", "KV Quantization", "KV Pruning (Token Eviction)"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ("압축 메커니즘", "모든 토큰의 KV를 낮은 비트(2/4bit)로 표현", "선택된 토큰의 KV를 제거하거나 merge"),
        ("메모리 절감 방식", "단위 토큰당 메모리 감소\n(2bit → 8배 압축)", "토큰 수 자체를 줄임\n(50% 제거 → 2배 압축)"),
        ("정보 보존", "모든 토큰 정보 (낮은 정밀도로) 보존", "선택된 토큰 정보만 full precision 보존"),
        ("하드웨어 친화성", "높음 (정수 연산 가속, CUDA 최적화 가능)", "중간 (irregular memory access 가능성)"),
        ("구현 복잡도", "중간 (역양자화 커널 필요)", "중간 (인덱싱 및 eviction 로직 필요)"),
        ("Semantic 제어", "불가 (모든 토큰 균등 압축)", "가능 (중요 토큰 선택적 보존)"),
        ("RAG 환경 적합성", "낮음 (키워드 토큰도 낮은 정밀도로 저하)", "높음 (RAG 키워드 토큰 full precision 유지)"),
        ("Long-context 확장", "좋음 (토큰 수 유지)", "매우 좋음 (토큰 수 감소로 attention 비용 절감)"),
        ("대표 방법", "KIVI, GEAR, QuIP, KVQuant", "H2O, SnapKV, PyramidKV, Scissorhands"),
        ("결합 가능성", "Pruning과 함께 사용 가능 (압축 토큰을 낮은 bit로)", "Quantization과 함께 사용 가능"),
        ("실효성 평가", "★★★★☆ (범용 서비스에 적합)", "★★★★★ (RAG/QA 특화 시 최우수)"),
    ]
    for row_data in comp_data:
        add_table_row(comp_table, row_data)

    doc.add_paragraph()

    set_heading(doc, "4.1 KV Pruning의 고유한 강점", 2)
    strengths = [
        ("Semantic Selectivity",
         "Quantization은 모든 토큰을 균등하게 압축하므로 중요한 토큰과 무관한 토큰을 "
         "구별하지 못한다. Pruning은 중요도 기반으로 토큰을 선별하므로 RAG 키워드, "
         "attention sink 등 semantically 중요한 토큰을 full precision으로 유지할 수 있다."),
        ("Attention Complexity 감소",
         "Quantization은 토큰 수를 줄이지 않으므로 attention 계산의 O(n²) 복잡도가 유지된다. "
         "Pruning은 실제 처리하는 KV 토큰 수를 줄여 attention 계산 비용 자체가 감소한다. "
         "이는 long-context 시나리오에서 latency 개선으로 직결된다."),
        ("Composability",
         "Pruning 이후 남은 토큰에 quantization을 추가 적용하면 두 방법의 장점을 결합할 수 있다. "
         "예: 중요 토큰은 full precision, 차선 토큰은 4bit, 나머지는 evict."),
        ("Interpretability",
         "어떤 토큰이 보존/제거되었는지 명시적으로 알 수 있어 디버깅과 분석이 용이하다. "
         "RAG 시나리오에서 검색 키워드가 보존되었는지 직접 확인 가능하다."),
        ("RAG-aware 확장 가능성",
         "외부 시스템(RAG retriever)의 신호를 pruning 결정에 자연스럽게 통합할 수 있다. "
         "Quantization은 외부 신호를 활용할 구조적 접점이 없으나, pruning은 "
         "importance score에 외부 신호를 추가 입력으로 사용하는 것이 자연스럽다."),
    ]
    for title, desc in strengths:
        p = doc.add_paragraph()
        p.add_run(f"▶ {title}: ").bold = True
        p.add_run(desc)

    set_heading(doc, "4.2 KV Pruning의 현재 한계", 2)
    limitations = [
        "Irreversibility: 제거된 토큰의 정보는 복구 불가능 (CaM, LESS 등에서 부분 해결 중)",
        "Importance Score 불안정성: Attention score 기반 중요도가 실제 output 기여도와 항상 일치하지 않음",
        "Query Dependency: Query-aware 방법은 query가 결정되어야 압축 가능, prefill cache 공유 어려움",
        "Hardware 비효율: Irregular memory pattern으로 인한 캐시 미스 및 메모리 대역폭 낭비 가능성",
    ]
    for lim in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(lim)

    doc.add_page_break()

    # ==========================================================================
    # 5. 나의 아이디어 평가
    # ==========================================================================
    set_heading(doc, "5. 나의 아이디어 평가: RAG-aware KV Pruning in vLLM", 1)

    set_heading(doc, "5.1 아이디어 요약", 2)
    doc.add_paragraph(
        "제안하는 시스템은 두 가지 핵심 요소로 구성된다:\n\n"
        "(1) vLLM KVShrinker 레이어: vLLM의 attention 레이어 전/후에 삽입되는 "
        "KV 압축 객체로, token importance score를 계산하여 budget 이하로 KV를 유지.\n\n"
        "(2) RAG-aware Importance Signal: BM25/TF-IDF 등 lexical retrieval이 "
        "반환하는 키워드 점수(term importance score)를 vLLM 추론 레이어로 전달하여, "
        "해당 키워드에 대응하는 토큰의 KV를 우선적으로 보존."
    )

    doc.add_paragraph()

    set_heading(doc, "5.2 아이디어 신규성(Novelty) 분석", 2)

    novelty_table = doc.add_table(rows=1, cols=3)
    novelty_table.style = 'Table Grid'
    for cell, text in zip(novelty_table.rows[0].cells, ["비교 대상", "차별점", "신규성 판단"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    novelty_data = [
        ("SnapKV\n(내부 query attention 활용)",
         "SnapKV는 LLM 내부의 instruction attention을 사용.\n본 아이디어는 외부 RAG retriever의 lexical score를 추가 신호로 활용.",
         "✅ 차별성 있음"),
        ("CompressKV\n(내부 retrieval head 활용)",
         "CompressKV는 모델 내부 retrieval attention head를 사용.\n본 아이디어는 모델 외부의 BM25/lexical 신호를 활용.",
         "✅ 차별성 있음"),
        ("HyperRAG\n(RAG KV 재사용)",
         "HyperRAG는 reranker KV를 generator가 재사용하는 것.\n본 아이디어는 검색 점수를 pruning 결정에 활용하는 것.",
         "✅ 명확히 다름"),
        ("ClusterKV\n(semantic clustering)",
         "ClusterKV는 모델 내부 KV semantic space에서 클러스터링.\n본 아이디어는 외부 lexical 신호로 중요 토큰 식별.",
         "✅ 접근법 상이"),
        ("H2O / Scissorhands\n(attention 기반 eviction)",
         "내부 attention score만 사용.\n본 아이디어는 내부 attention + 외부 RAG 신호의 결합.",
         "✅ 추가 신호 활용"),
        ("기존 RAG 최적화 연구",
         "RAG의 검색 결과를 LLM 입력으로 사용하는 방법은 많지만,\n검색 신호를 KV pruning에 직접 활용하는 연구는 없음.",
         "✅ 미개척 영역"),
    ]
    for row_data in novelty_data:
        add_table_row(novelty_table, row_data)

    doc.add_paragraph()

    set_heading(doc, "5.3 실효성 평가", 2)

    doc.add_paragraph("■ 강점", style='Heading 3')
    pros = [
        ("원칙적 타당성",
         "BM25 등 lexical 검색이 반환하는 키워드는 사용자의 질문과 가장 관련 있는 토큰이다. "
         "이 토큰들의 KV를 보존하는 것은 retrieved knowledge를 올바르게 활용하기 위한 "
         "필요조건에 해당한다. 이는 attention score 기반 방법이 놓칠 수 있는 "
         "'semantically important but attention-sparse' 토큰을 커버한다."),
        ("RAG 파이프라인 호환성",
         "현재 production RAG 시스템(LangChain, LlamaIndex, vLLM + RAG)은 "
         "이미 lexical retrieval 점수를 생성한다. 추가적인 모델 재학습 없이 "
         "기존 RAG 파이프라인에 KVShrinker를 삽입하는 방식으로 구현 가능."),
        ("vLLM 통합 가능성",
         "vLLM은 PagedAttention 기반으로 블록 단위 KV 관리를 한다. "
         "KVShrinker를 attention layer 이후 hook 또는 별도 레이어로 구현하면 "
         "기존 vLLM 코드베이스의 최소 수정으로 통합 가능."),
        ("성능 하락 최소화 가설",
         "RAG QA 태스크에서 답변 생성에 필요한 토큰은 검색된 문서의 핵심 키워드들이다. "
         "이들을 보존함으로써 관련 없는 토큰을 공격적으로 제거해도 품질 저하가 "
         "최소화될 것이라는 가설은 실험으로 검증 가능하다."),
    ]
    for title, desc in pros:
        p = doc.add_paragraph()
        p.add_run(f"  ▶ {title}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph("■ 잠재적 한계 및 대응 방안", style='Heading 3')
    cons = [
        ("Lexical ≠ Semantic 괴리",
         "BM25가 중요하다고 판단한 키워드와 LLM이 실제로 attention하는 토큰이 다를 수 있다.",
         "대응: Dense retrieval 점수(embedding similarity)를 추가 신호로 결합하거나, "
         "내부 attention score와 RAG score를 linear combination으로 통합."),
        ("토큰화(Tokenization) 불일치",
         "BM25는 단어 단위로 동작하지만 LLM은 sub-word 토큰을 사용한다.",
         "대응: Sub-word 토큰을 원래 단어로 re-aggregate한 후 BM25 점수를 매핑. "
         "예: 'retriev-al' → 두 토큰 모두 'retrieval'의 BM25 점수 상속."),
        ("Multi-hop Reasoning에서의 불안정성",
         "Complex reasoning에서는 RAG 키워드 외의 중간 추론 단계 토큰도 중요할 수 있다.",
         "대응: Chain-of-thought(CoT) 생성 단계에서는 pruning 강도를 완화하는 "
         "adaptive strategy 적용 (Lethe 방법론 참조)."),
        ("RAG-LLM 인터페이스 설계",
         "RAG retriever와 vLLM inference 사이의 keyword score 전달 인터페이스가 필요하다.",
         "대응: vLLM의 SamplingParams 또는 별도 metadata 필드를 통해 "
         "keyword mask/score를 전달하는 API 설계."),
    ]
    for title, limit, solution in cons:
        p = doc.add_paragraph()
        p.add_run(f"  △ {title}: ").bold = True
        p.add_run(f"{limit}\n  → {solution}")

    doc.add_paragraph()

    set_heading(doc, "5.4 vLLM 구현 설계", 2)
    doc.add_paragraph(
        "구현 계획은 다음 단계로 진행한다:\n\n"
        "1단계: vLLM fork 생성 및 환경 설정\n"
        "   - github.com/vllm-project/vllm를 fork\n"
        "   - vllm/attention/ 하위에 kv_shrinker.py 모듈 추가\n\n"
        "2단계: KVShrinker 클래스 구현\n"
        "   - class KVShrinker: 입력 = KV tensor, RAG keyword scores, budget ratio\n"
        "   - 중요도 계산: attention_score × (1 + α × rag_score)\n"
        "   - Eviction: top-k 보존, 나머지 제거 또는 merge\n\n"
        "3단계: PagedAttention 통합\n"
        "   - Block 단위 eviction을 PagedAttention block table과 동기화\n"
        "   - 블록 내 토큰을 개별 evict하거나 블록 전체를 release하는 선택 구현\n\n"
        "4단계: RAG 인터페이스 구현\n"
        "   - SamplingParams에 rag_keyword_scores: Dict[str, float] 필드 추가\n"
        "   - Tokenization 후 token → keyword score 매핑 함수 구현\n\n"
        "5단계: vLLM Pull Request 준비\n"
        "   - 기존 KV 압축 방법들과 동일한 API 인터페이스로 구현\n"
        "   - 벤치마크 결과, 단위 테스트, 문서화 포함"
    )

    doc.add_page_break()

    # ==========================================================================
    # 6. 미개척 연구 방향
    # ==========================================================================
    set_heading(doc, "6. 미개척 연구 방향 및 Top-tier 논문 아이디어", 1)

    set_heading(doc, "6.1 현재 미개척 영역 분석", 2)
    gaps = [
        ("Cross-system KV Signal Integration",
         "내/외부 신호 결합",
         "RAG retriever, reranker, 내부 retrieval head 등 여러 소스의 중요도 신호를 통합하여 KV pruning을 수행하는 연구 없음. Multi-source importance aggregation이 미개척."),
        ("RAG-aware KV Pruning",
         "본 아이디어 (미개척)",
         "RAG 파이프라인의 lexical/dense retrieval 점수를 vLLM KV pruning에 직접 활용하는 연구 없음. RAG + KV pruning의 end-to-end 최적화 연구 부재."),
        ("Token-level + Dimension-level 2D Compression",
         "직교 압축 결합",
         "Token eviction(H2O, SnapKV)과 dimension pruning(ThinK, SparK)을 동시에 최적화하는 통합 프레임워크 없음. 두 차원을 joint하게 최적화하면 더 높은 압축률 달성 가능."),
        ("Theoretical Optimality under RAG Constraint",
         "이론 기여",
         "RAG 환경에서 특정 토큰을 반드시 보존해야 한다는 제약 조건 하의 optimal KV pruning 이론 없음. OBCache의 OBS 프레임워크를 constrained optimization으로 확장 가능."),
        ("Online RAG-KV Co-optimization",
         "공동 최적화",
         "RAG retrieval과 KV compression을 분리된 파이프라인이 아닌 공동으로 최적화하는 연구 없음. Retriever가 KV-friendly한 결과를 반환하도록 학습하는 연구."),
        ("Speculative KV Pruning",
         "투기적 압축",
         "Speculative decoding과 KV pruning을 결합하여, draft model의 KV를 target model에 재사용하면서 동시에 pruning하는 방법 연구 부재."),
    ]

    for title, tag, desc in gaps:
        p = doc.add_paragraph()
        p.add_run(f"[{tag}] {title}\n").bold = True
        p.add_run(f"  → {desc}")

    set_heading(doc, "6.2 NeurIPS/EMNLP 타겟 아이디어 (메인 제안)", 2)

    set_heading(doc, "◆ 아이디어 1: RAG-KV — Retrieval-Augmented KV Cache Pruning", 3)
    doc.add_paragraph(
        "타겟 학회: NeurIPS 2026 또는 EMNLP 2026\n\n"
        "핵심 아이디어:\n"
        "RAG 파이프라인에서 lexical retriever(BM25)와 dense retriever(DPR/ColBERT)가 "
        "생성하는 토큰 중요도 신호를 vLLM KV 캐시 pruning의 guidance signal로 활용. "
        "새로운 Retrieval-Guided Token Importance Score(RTIS)를 정의:\n\n"
        "    RTIS(t) = α · AttentionScore(t) + β · LexicalScore(t) + γ · DenseScore(t)\n\n"
        "여기서 LexicalScore(t)는 토큰 t가 BM25 top-k keyword에 속하는 경우의 점수이고, "
        "DenseScore(t)는 dense retrieval embedding과의 유사도 기반 점수.\n\n"
        "뚜렷한 기여(Contributions):\n"
        "  1. RAG-aware KV pruning의 최초 공식화 및 이론적 동기 부여\n"
        "  2. 다양한 retriever 유형(lexical, dense, hybrid)과의 통합 프레임워크\n"
        "  3. vLLM PagedAttention 통합 구현 (오픈소스)\n"
        "  4. RAG QA, Long-doc QA, Multi-hop QA 등에서의 체계적 평가"
    )

    set_heading(doc, "◆ 아이디어 2: KeyGuard — 2D KV Compression with RAG-aware Token and Channel Pruning", 3)
    doc.add_paragraph(
        "타겟 학회: NeurIPS 2026\n\n"
        "핵심 아이디어:\n"
        "Token-level pruning(어떤 토큰을 제거할지)과 dimension-level pruning(각 KV의 "
        "어떤 채널을 제거할지)을 동시에 최적화하는 통합 프레임워크. "
        "RAG 키워드 토큰은 full dimension 보존, 비-키워드 토큰은 dimension도 함께 압축.\n\n"
        "    KV_compressed(t) = \n"
        "      Full KV(t)           if t ∈ RAG_keywords ∪ AttentionSink\n"
        "      KV_dimPruned(t)      if t ∈ TopAttention \\ RAG_keywords  \n"
        "      Evict(t)             otherwise\n\n"
        "뚜렷한 기여:\n"
        "  1. Token-level + Dimension-level의 최초 통합 최적화\n"
        "  2. RAG 신호를 two-dimensional 압축 결정의 hard constraint로 활용\n"
        "  3. 동일 메모리 budget에서 기존 방법보다 ~15% 성능 향상 목표"
    )

    set_heading(doc, "◆ 아이디어 3: AdaptKV-RAG — Adaptive Layer-wise RAG-guided KV Budget Allocation", 3)
    doc.add_paragraph(
        "타겟 학회: EMNLP 2026\n\n"
        "핵심 아이디어:\n"
        "PyramidKV의 layer-wise budget allocation을 RAG 신호로 guide하는 방법. "
        "RAG retrieved document가 처리되는 특정 layer에서 RAG 키워드 토큰의 "
        "attention entropy를 모니터링하고, entropy가 높은 layer(RAG 내용이 "
        "광범위하게 처리되는 layer)에 더 많은 budget을 동적으로 할당.\n\n"
        "혁신 포인트: '언제 어느 layer에서 RAG 키워드가 중요한가'를 자동으로 감지하는 "
        "layer-adaptive RAG signal weighting 메커니즘."
    )

    doc.add_page_break()

    # ==========================================================================
    # 7. 실험 방법론
    # ==========================================================================
    set_heading(doc, "7. 실험 방법론", 1)

    set_heading(doc, "7.1 실험 환경", 2)
    doc.add_paragraph(
        "하드웨어: 다수 GPU 보유 환경 (권장: NVIDIA A100 80GB × 4 이상)\n"
        "소프트웨어: vLLM (최신 버전), PyTorch 2.x, CUDA 12.x\n"
        "기반 모델: Llama-3.1-8B, Llama-3.1-70B, Mistral-7B-v0.3\n"
        "(다양한 크기에서 scalability 검증 필요)"
    )

    set_heading(doc, "7.2 평가 데이터셋", 2)
    datasets = [
        ("NaturalQuestions (NQ)", "RAG QA", "RAG 키워드 보존 효과 측정에 핵심"),
        ("TriviaQA", "RAG QA", "다양한 도메인의 사실 기반 QA"),
        ("HotpotQA", "Multi-hop RAG QA", "복잡한 reasoning에서의 pruning 안정성 검증"),
        ("LongBench", "Long-context 이해", "긴 문서 처리 성능 검증"),
        ("RULER", "Long-context 합성 데이터", "KV 압축률 vs 성능 trade-off 정밀 측정"),
        ("SCROLLS (Qasper, NarrativeQA)", "문서 기반 QA", "실제 RAG 시나리오와 유사"),
    ]
    ds_table = doc.add_table(rows=1, cols=3)
    ds_table.style = 'Table Grid'
    for cell, text in zip(ds_table.rows[0].cells, ["데이터셋", "유형", "선택 이유"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_data in datasets:
        add_table_row(ds_table, row_data)

    doc.add_paragraph()

    set_heading(doc, "7.3 비교 베이스라인", 2)
    baselines = [
        "Full KV (압축 없음) — 상한선(oracle)",
        "H2O — 가장 기본적인 attention score 기반 eviction",
        "SnapKV — query-aware SOTA (NeurIPS 2024)",
        "PyramidKV — layer-wise SOTA",
        "KIVI (2-bit quantization) — quantization 대표",
        "GEAR — quantization+low-rank 대표",
        "Proposed: RAG-KV (본 아이디어) — RAG-aware pruning",
        "Ablation: RAG-KV w/o lexical signal — RAG 신호의 기여도 측정",
        "Ablation: RAG-KV w/o dense signal — dense vs lexical 비교",
    ]
    for b in baselines:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b)

    set_heading(doc, "7.4 평가 지표", 2)
    metrics = [
        ("Exact Match (EM) / F1", "QA 정확도", "핵심 성능 지표"),
        ("Perplexity (PPL)", "언어 모델링 품질", "압축으로 인한 품질 저하 측정"),
        ("KV Cache Memory (GB)", "메모리 절감", "절대적 메모리 사용량"),
        ("KV Compression Ratio", "압축률", "원본 대비 KV 크기 비율"),
        ("Throughput (token/s)", "처리량", "서비스 성능 지표"),
        ("Time-to-First-Token (TTFT)", "지연시간", "latency 지표"),
        ("RAG Keyword Retention Rate", "키워드 보존률", "RAG 신호 활용 검증용 지표 (신규)"),
    ]
    metric_table = doc.add_table(rows=1, cols=3)
    metric_table.style = 'Table Grid'
    for cell, text in zip(metric_table.rows[0].cells, ["지표", "측정 대상", "비고"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_data in metrics:
        add_table_row(metric_table, row_data)

    doc.add_paragraph()

    set_heading(doc, "7.5 핵심 실험 설계", 2)
    experiments = [
        ("Exp 1: 압축률-성능 Trade-off 곡선",
         "KV budget을 10%–90%로 변화시키며 EM/F1 변화 측정. "
         "RAG-KV가 동일 budget에서 SnapKV, H2O, PyramidKV 대비 일관되게 우수함을 보임."),
        ("Exp 2: RAG 키워드 중요성 검증",
         "RAG 키워드 토큰을 (a) 보존, (b) 우선 제거했을 때 성능 비교. "
         "키워드 보존이 성능에 미치는 영향을 정량화."),
        ("Exp 3: Lexical vs Dense Signal 비교",
         "BM25 단독, DPR 단독, BM25+DPR 결합 신호의 성능 비교. "
         "최적 신호 조합 방법 탐색."),
        ("Exp 4: 모델 크기 Scalability",
         "7B, 13B, 70B 모델에서 RAG-KV의 성능/메모리 확장성 검증."),
        ("Exp 5: RAG 문서 수 vs KV Budget",
         "검색된 문서가 1개, 3개, 5개, 10개일 때 최적 budget 전략 분석."),
        ("Exp 6: vLLM 서빙 벤치마크",
         "동일 GPU에서 RAG-KV 적용 전/후의 throughput, latency, batch size 비교."),
    ]
    for title, desc in experiments:
        p = doc.add_paragraph()
        p.add_run(f"  ▶ {title}\n").bold = True
        p.add_run(f"    {desc}")

    doc.add_page_break()

    # ==========================================================================
    # 8. 결론
    # ==========================================================================
    set_heading(doc, "8. 결론 및 향후 계획", 1)

    doc.add_paragraph(
        "본 보고서는 2022–2025년 KV 캐시 압축 연구 27편을 분석하고, "
        "제안 아이디어인 RAG-aware KV Pruning의 실현 가능성을 평가하였다."
    )

    set_heading(doc, "8.1 핵심 결론", 2)
    conclusions = [
        "KV Pruning은 Quantization 대비 RAG 환경에서 semantic selectivity와 "
        "attention complexity 감소라는 명확한 강점을 가진다.",
        "제안 아이디어(RAG-aware KV Pruning)는 기존 연구에서 다루지 않은 "
        "미개척 영역으로, NeurIPS/EMNLP 수준의 novelty를 갖춘다.",
        "vLLM PagedAttention과의 통합이 기술적으로 가능하며, "
        "최소한의 코드 수정으로 구현 가능하다.",
        "Lexical+dense 신호 결합, 2D 압축(token+dimension), "
        "이론적 최적성 분석이 논문의 핵심 기여가 될 수 있다.",
    ]
    for c in conclusions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(c)

    set_heading(doc, "8.2 단계별 로드맵", 2)
    roadmap = [
        ("Phase 1 (1–2개월)", "vLLM fork, KVShrinker 기본 구현, SnapKV/H2O 재현"),
        ("Phase 2 (2–3개월)", "RAG interface 구현, BM25 신호 통합, 초기 실험"),
        ("Phase 3 (3–4개월)", "Dense signal 통합, 주요 실험 완료, ablation study"),
        ("Phase 4 (4–5개월)", "논문 작성, vLLM PR 준비, NeurIPS/EMNLP submission"),
    ]
    rm_table = doc.add_table(rows=1, cols=2)
    rm_table.style = 'Table Grid'
    for cell, text in zip(rm_table.rows[0].cells, ["단계", "주요 작업"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_data in roadmap:
        add_table_row(rm_table, row_data)

    # ==========================================================================
    # 저장
    # ==========================================================================
    output_path = os.path.join(os.path.dirname(__file__), "KV_Cache_Pruning_Report.docx")
    doc.save(output_path)
    print(f"  보고서 저장 완료: {output_path}")
    return output_path


if __name__ == "__main__":
    download_papers()
    output = generate_report()
    print(f"\n✅ 완료! 보고서: {output}")
    print(f"   논문 폴더: {PAPERS_DIR}")
